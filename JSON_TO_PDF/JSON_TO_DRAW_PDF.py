import json
import re
import shutil
import subprocess
from pathlib import Path
from copy import deepcopy
from typing import Any, Dict, Union, Optional

import pikepdf
import fitz  # PyMuPDF
from lxml import etree as ET
try:
    from docx import Document
except ImportError:
    Document = None

# ============================================================
# Paths
# ============================================================
BASE_DIR = Path(__file__).resolve().parent

PDF_IN = BASE_DIR / "dd2977.pdf"
DOCX_IN = Path(__file__).parent / "dd2977.docx"
PDF_OUT = BASE_DIR / "dd2977_filled.pdf"
JSON_IN = BASE_DIR / "input_draw.json"

# ============================================================
# XFA namespace
# ============================================================
XFA_DATA_NS = "http://www.xfa.org/schema/xfa-data/1.0/"
NSMAP = {"xfa": XFA_DATA_NS}

# ============================================================
# Utility: Clean text
# ============================================================
def clean_ascii(s: str):
    """
    Remove zero-width spaces, Adobe special NBSP, en-space, em-space, 
    and any unicode outside the standard printable ASCII range.
    """
    if s is None:
        return ""
    return re.sub(r"[^\x20-\x7E\n\r\t]", "", str(s))

# ============================================================
# XFA Logic
# ============================================================
def find_xfa_datasets(pdf: pikepdf.Pdf):
    acroform = pdf.Root.get("/AcroForm", None)
    if acroform is None:
        raise RuntimeError("No /AcroForm in PDF")

    xfa = acroform.get("/XFA", None)
    if xfa is None:
        raise RuntimeError("No /XFA array found")

    # XFA is an array of [key, stream, key, stream, ...]
    for i in range(0, len(xfa), 2):
        if str(xfa[i]) == "datasets":
            return i + 1, xfa[i + 1].read_bytes()

    raise RuntimeError("datasets XFA packet not found")

def rebuild_datasets_in_place(xml_root, data):
    # 🔹 Navigate to main nodes (all already exist in the template)
    data_node = xml_root.find("xfa:data", NSMAP)
    if data_node is None:
        # Fallback if namespace prefix is missing or different
        data_node = xml_root.find("{http://www.xfa.org/schema/xfa-data/1.0/}data")
    
    if data_node is None:
        # Try finding without namespace if strictly necessary, but XFA usually has it.
        # Let's print root children if we fail?
        pass

    form1 = data_node.find("form1")
    page1 = form1.find("Page1")

    # Simple fields mapping based on the original file content
    one = page1.find("One")
    two = page1.find("Two")
    A = page1.find("A"); B = page1.find("B"); C = page1.find("C")
    D = page1.find("D"); E = page1.find("E"); F = page1.find("F")
    G = page1.find("G"); H = page1.find("H")

    eleven = page1.find("Eleven")
    ten = page1.find("Ten")
    twelve = page1.find("Twelve")
    part4 = page1.find("Part4thru9")

    # ============================================================
    # Fill Page1 fields
    # ============================================================
    if one is not None:
        one.text = clean_ascii(data.get("mission_task_and_description", ""))

    raw_date = data.get("date", "")
    clean_date = clean_ascii(raw_date).replace("-", "")
    if two is not None:
        two.text = clean_ascii(clean_date)

    prep = data.get("prepared_by", {}) or {}
    if A is not None: A.text = clean_ascii(prep.get("name_last_first_middle_initial", ""))
    if B is not None: B.text = clean_ascii(prep.get("rank_grade", ""))
    if C is not None: C.text = clean_ascii(prep.get("duty_title_position", ""))
    if D is not None: D.text = clean_ascii(prep.get("unit", ""))
    if E is not None: E.text = clean_ascii(prep.get("work_email", ""))
    if F is not None: F.text = clean_ascii(prep.get("telephone", ""))
    if G is not None: G.text = clean_ascii(prep.get("uic_cin", ""))
    if H is not None: H.text = clean_ascii(prep.get("training_support_or_lesson_plan_or_opord", ""))

    if eleven is not None:
        eleven.text = clean_ascii(data.get("overall_supervision_plan", ""))

    # ============================================================
    # Block 10 — Overall RRL
    # ============================================================
    if ten is not None:
        overall = clean_ascii((data.get("overall_residual_risk_level") or "").upper())

        # Reset all
        for tag in ["EHigh", "High", "Med", "Low"]:
            node = ten.find(tag)
            if node is not None:
                node.text = "0"

        map_rrl = {"EH": "EHigh", "H": "High", "M": "Med", "L": "Low"}
        if overall in map_rrl:
            tgt = ten.find(map_rrl[overall])
            if tgt is not None:
                tgt.text = "1"

    # ============================================================
    # Block 12 — Approval / Disapproval
    # ============================================================
    if twelve is not None:
        appr = data.get("approval_or_disapproval_of_mission_or_task", {}) or {}

        approve = twelve.find("Approve")
        dis = twelve.find("Disapprove")

        if approve is not None:
            approve.text = "1" if appr.get("approve") else "0"

        if dis is not None:
            dis.text = "1" if appr.get("disapprove") else "0"

    # ============================================================
    # Hazard Table (Part4thru9)
    # ============================================================
    if part4 is not None:
        template_row = part4.find("Row1")
        
        if template_row is not None:
            # Remove existing rows (keep template in memory)
            # Note: In XFA, repeated elements are usually siblings.
            # We need to be careful not to remove the only Row1 if we need it for cloning.
            # But here we clone it first.
            
            # Find all Row1 elements and remove them
            for child in list(part4):
                if child.tag == "Row1":
                    part4.remove(child)

            # Rebuild rows
            for st in data.get("subtasks", []):
                row = deepcopy(template_row)

                # Subtask
                sub = row.find("Subtask-Substep")
                if sub is not None:
                    sub.text = clean_ascii((st.get("subtask") or {}).get("name", ""))

                # Hazard
                haz = row.find("Hazard")
                if haz is not None:
                    haz.text = clean_ascii(st.get("hazard", ""))

                # Initial Risk Level
                irl = row.find("InitialRiskLevel")
                if irl is not None:
                    irl.text = clean_ascii((st.get("initial_risk_level") or "").upper())

                # Control
                ctrl = row.find("Control")
                if ctrl is not None:
                    ctrl.text = clean_ascii("\n".join((st.get("control") or {}).get("values", [])))

                # HOW / WHO
                table2 = row.find("Table2")
                if table2 is not None:
                    r1 = table2.find("Row1"); tf1 = r1.find("TextField1") if r1 is not None else None
                    r2 = table2.find("Row2"); tf2 = r2.find("TextField2") if r2 is not None else None

                    how_vals = (st.get("how_to_implement") or {}).get("how", {}).get("values", [])
                    who_vals = (st.get("how_to_implement") or {}).get("who", {}).get("values", [])

                    if tf1 is not None:
                        tf1.text = clean_ascii("\n".join(how_vals))
                    if tf2 is not None:
                        tf2.text = clean_ascii("\n".join(who_vals))

                # Residual Risk Level
                rrl = row.find("RRL")
                if rrl is not None:
                    rrl.text = clean_ascii((st.get("residual_risk_level") or "").upper())

                part4.append(row)

# ============================================================
# Main Exported Function
# ============================================================
def generate_draw_pdf(data: Dict[str, Any], output_path: Union[str, Path]):
    """
    Generates a filled DD2977 PDF based on the provided data using XFA injection.
    """
    output_path = Path(output_path)
    
    if not PDF_IN.exists():
        raise FileNotFoundError(f"Template PDF not found at {PDF_IN}")

    pdf = pikepdf.Pdf.open(PDF_IN)
    try:
        xfa_index, datasets_bytes = find_xfa_datasets(pdf)

        xml_root = ET.fromstring(datasets_bytes)
        rebuild_datasets_in_place(xml_root, data)

        new_xml = ET.tostring(xml_root, encoding="utf-8", xml_declaration=False)

        # Update the XFA stream
        datasets_stream = pdf.Root["/AcroForm"]["/XFA"][xfa_index]
        
        # pikepdf stream update
        datasets_stream.write(new_xml)

        pdf.save(output_path)
        print(f"✅ SUCCESS — Filled (XFA) DD2977 created at: {output_path}")
    finally:
        pdf.close()

def fill_docx_template(docx_path: Path, output_pdf_path: Path, data: Dict[str, Any]):
    if Document is None:
        raise RuntimeError("python-docx not installed")
        
    doc = Document(docx_path)
    table = doc.tables[0]
    
    def set_cell(r, c, text, label=None):
        try:
            cell = table.rows[r].cells[c]
            # If label=True, we append to existing text (assuming label is there)
            # For this specific form, we just append a newline + value
            if label and cell.paragraphs:
                p = cell.paragraphs[-1]
                run = p.add_run(f" {text}")
                run.bold = True
            else:
                cell.text = text
        except IndexError:
            pass # Row/Col doesn't exist

    def insert_row_after(ref_row_idx, row_to_copy_idx):
        """
        Copy the row at `row_to_copy_idx` and insert it after `ref_row_idx`.
        """
        ref_row = table.rows[ref_row_idx]
        copy_src = table.rows[row_to_copy_idx]
        new_xml = deepcopy(copy_src._element)
        ref_row._element.addnext(new_xml)

    # Header
    set_cell(1, 0, data.get("mission_task_and_description", ""), label=True)
    set_cell(1, 13, data.get("date", ""), label=True)
    
    prep = data.get("prepared_by") or {}
    set_cell(3, 0, prep.get("name_last_first_middle_initial", ""), label=True)
    set_cell(3, 7, prep.get("rank_grade", ""), label=True)
    set_cell(3, 11, prep.get("duty_title_position", ""), label=True)
    set_cell(4, 0, prep.get("unit", ""), label=True)
    set_cell(4, 2, prep.get("work_email", ""), label=True)
    set_cell(4, 9, prep.get("telephone", ""), label=True)
    set_cell(5, 0, prep.get("uic_cin", ""), label=True)
    set_cell(5, 2, prep.get("training_support_or_lesson_plan_or_opord", ""), label=True)

    # Risks (Dynamic Rows)
    subtasks = data.get("subtasks") or []
    
    # Template rows are at index 8 (Main) and 9 (Who)
    # We track the index of the last used "Who" row
    last_who_row_idx = 9
    
    for i, st in enumerate(subtasks):
        if i == 0:
            # Use the existing template rows
            curr_main_idx = 8
            curr_who_idx = 9
        else:
            # Insert new pair of rows after the last "Who" row
            # We insert them in reverse order so they end up: Main, Who
            
            # 1. Insert copy of Row 9 (Who) after last_who_row_idx
            insert_row_after(last_who_row_idx, 9)
            # 2. Insert copy of Row 8 (Main) after last_who_row_idx
            insert_row_after(last_who_row_idx, 8)
            
            # Update indices
            curr_main_idx = last_who_row_idx + 1
            curr_who_idx = last_who_row_idx + 2
            last_who_row_idx += 2

        # Fill the rows
        subtask_info = st.get('subtask') or {}
        set_cell(curr_main_idx, 1, subtask_info.get('name', ''))
        set_cell(curr_main_idx, 3, st.get('hazard', ''))
        set_cell(curr_main_idx, 5, st.get('initial_risk_level', ''))
        
        control_info = st.get("control") or {}
        ctrls = control_info.get("values", [])
        ctrl_text = "\n".join(ctrls) if isinstance(ctrls, list) else str(ctrls)
        set_cell(curr_main_idx, 8, ctrl_text)
        
        how_impl = st.get("how_to_implement") or {}
        how_info = how_impl.get("how") or {}
        how = how_info.get("values", [])
        how_text = "\n".join(how) if isinstance(how, list) else str(how)
        set_cell(curr_main_idx, 12, how_text)
        
        who_info = how_impl.get("who") or {}
        who = who_info.get("values", [])
        who_text = "\n".join(who) if isinstance(who, list) else str(who)
        set_cell(curr_who_idx, 12, who_text)
        
        set_cell(curr_main_idx, 14, st.get('residual_risk_level', ''))

    # Overall Risk (Find the row index, as it has shifted)
    # It was originally Row 10. It is now after the last_who_row_idx.
    overall_row_idx = last_who_row_idx + 1
    overall = data.get("overall_residual_risk_level", "").upper()
    set_cell(overall_row_idx, 0, f"  {overall}", label=True)

    # Save to temp docx
    temp_docx = output_pdf_path.with_suffix(".temp.docx")
    doc.save(temp_docx)
    
    # Convert to PDF using LibreOffice
    # Assumes 'soffice' is on PATH (which api_server checks)
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(output_pdf_path.parent), str(temp_docx)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        # LibreOffice saves as temp.pdf
        temp_pdf = output_pdf_path.with_suffix(".temp.pdf")
        if temp_pdf.exists():
            shutil.move(temp_pdf, output_pdf_path)
        else:
            # Fallback if name is different
            pass
    except Exception as e:
        print(f"DOCX conversion failed: {e}")
    finally:
        if temp_docx.exists():
            temp_docx.unlink()

def render_preview_pdf(input_path: Union[str, Path], output_path: Union[str, Path], data: Optional[Dict[str, Any]] = None):
    """
    Creates a preview version of the PDF.
    If DOCX_IN exists, fills it and converts to PDF.
    Else if 'data' is provided, generates a summary PDF using PyMuPDF.
    Else, falls back to copying the XFA file.
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    # Try DOCX template first
    if DOCX_IN.exists() and data and Document:
        try:
            fill_docx_template(DOCX_IN, output_path, data)
            if output_path.exists():
                return
        except Exception as e:
            print(f"DOCX template filling failed: {e}")

    if not data:
        shutil.copy2(input_path, output_path)
        return

    # Generate a summary PDF (Fallback)
    doc = fitz.open()
    page = doc.new_page()
    
    # Helper to write text
    y = 50
    line_height = 14
    margin = 50
    
    def write_line(text, size=11, bold=False):
        nonlocal y, page
        # Use default font to avoid loading issues
        page.insert_text((margin, y), str(text), fontsize=size)
        y += line_height
        if y > 800:
            page = doc.new_page()
            y = 50

    write_line("DD2977 CONTENT PREVIEW", size=16, bold=True)
    y += 10
    write_line("(This is a generated summary. Download the file to view the official XFA form.)", size=10)
    y += 20

    # Header Info
    write_line(f"Mission/Task: {data.get('mission_task_and_description', 'N/A')}", bold=True)
    write_line(f"Date: {data.get('date', 'N/A')}")
    
    prep = data.get("prepared_by", {})
    write_line(f"Prepared By: {prep.get('name_last_first_middle_initial', 'N/A')} ({prep.get('rank_grade', '')})")
    write_line(f"Unit: {prep.get('unit', 'N/A')}")
    y += 10

    # Risk Assessment
    write_line("RISK ASSESSMENT:", size=12, bold=True)
    y += 5

    for i, subtask in enumerate(data.get("subtasks", []), 1):
        st_name = (subtask.get("subtask") or {}).get("name", "Unknown Subtask")
        write_line(f"{i}. Subtask: {st_name}", bold=True)
        
        haz = subtask.get("hazard", "N/A")
        write_line(f"   Hazard: {haz}")
        
        ctrl = (subtask.get("control") or {}).get("values", [])
        ctrl_text = "; ".join(ctrl) if isinstance(ctrl, list) else str(ctrl)
        write_line(f"   Controls: {ctrl_text}")
        
        risk = subtask.get("residual_risk_level", "N/A")
        write_line(f"   Residual Risk: {risk}")
        y += 10

    # Overall Risk
    overall = data.get("overall_residual_risk_level", "N/A")
    write_line(f"OVERALL RESIDUAL RISK LEVEL: {overall}", size=12, bold=True)

    doc.save(output_path)
    doc.close()

# ============================================================
# CLI Entrypoint
# ============================================================
def main():
    def load_json():
        with JSON_IN.open("r", encoding="utf-8") as f:
            return json.load(f)
            
    data = load_json()
    generate_draw_pdf(data, PDF_OUT)

if __name__ == "__main__":
    main()
